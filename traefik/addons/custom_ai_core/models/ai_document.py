"""RAG knowledge base — documents, chunking, and vector search."""

import json
import logging

from odoo import api, fields, models

_logger = logging.getLogger(__name__)

_PGVECTOR_AVAILABLE = None  # cached after first check


def _check_pgvector(env) -> bool:
    global _PGVECTOR_AVAILABLE
    if _PGVECTOR_AVAILABLE is None:
        try:
            env.cr.execute(
                "SELECT 1 FROM pg_available_extensions WHERE name='vector' AND installed_version IS NOT NULL"
            )
            _PGVECTOR_AVAILABLE = bool(env.cr.fetchone())
        except Exception:
            _PGVECTOR_AVAILABLE = False
    return _PGVECTOR_AVAILABLE


class AiDocument(models.Model):
    _name = "ai.document"
    _description = "AI Knowledge Base Document"
    _inherit = ["mail.thread"]
    _order = "name"

    name = fields.Char(required=True)
    company_id = fields.Many2one(
        "res.company", required=True, default=lambda s: s.env.company, index=True
    )

    # Content — either plain text or an uploaded file
    content = fields.Text("Plain Text Content")
    attachment_id = fields.Many2one("ir.attachment", "Uploaded File", ondelete="set null")
    mimetype = fields.Char(related="attachment_id.mimetype", store=True)

    # Status
    status = fields.Selection(
        [
            ("draft", "Draft"),
            ("pending", "Pending Indexing"),
            ("indexed", "Indexed"),
            ("error", "Error"),
        ],
        default="draft",
        tracking=True,
    )
    chunk_count = fields.Integer(compute="_compute_chunk_count", store=True)
    last_indexed = fields.Datetime(readonly=True)
    index_error = fields.Text(readonly=True)

    # Metadata
    source_url = fields.Char("Source URL")
    tags = fields.Char("Tags (comma-separated)")
    is_active = fields.Boolean(default=True)

    chunk_ids = fields.One2many("ai.document.chunk", "document_id", "Chunks")

    @api.depends("chunk_ids")
    def _compute_chunk_count(self):
        for rec in self:
            rec.chunk_count = len(rec.chunk_ids)

    def action_index(self):
        """Chunk and embed this document. Offloads to Celery when available."""
        from odoo.addons.custom_ai_core.lib.task_bridge import dispatch  # noqa: PLC0415

        for doc in self:
            dispatched = dispatch(
                "tasks.ai_tasks.embed_document",
                args=(self.env.cr.dbname, doc.id),
            )
            if not dispatched:
                doc._do_index()

    def action_delete_index(self):
        """Remove all chunks/embeddings for this document."""
        for doc in self:
            doc.chunk_ids.unlink()
            doc.write({"status": "draft", "last_indexed": False, "index_error": False})

    def action_view_chunks(self):
        """Smart-button: open the chunks belonging to this document."""
        self.ensure_one()
        return {
            "type": "ir.actions.act_window",
            "name": f"Chunks — {self.name}",
            "res_model": "ai.document.chunk",
            "view_mode": "list,form",
            "domain": [("document_id", "=", self.id)],
            "context": {"default_document_id": self.id},
            "target": "current",
        }

    def _do_index(self):
        self.ensure_one()
        self.write({"status": "pending"})
        try:
            text = self._extract_text()
            if not text:
                self.write({"status": "error", "index_error": "No text content found."})
                return

            chunks = self._chunk_text(text)
            provider = self.env["ai.provider"].get_default_provider()
            instance = provider.get_provider_instance() if provider else None

            # Remove old chunks
            self.chunk_ids.unlink()

            for idx, chunk_text in enumerate(chunks):
                embedding = []
                if instance:
                    try:
                        embedding = instance.embed(chunk_text)
                    except Exception as exc:
                        _logger.warning("Embedding failed for chunk %d: %s", idx, exc)

                self.env["ai.document.chunk"].create(
                    {
                        "document_id": self.id,
                        "company_id": self.company_id.id,
                        "chunk_index": idx,
                        "content": chunk_text,
                        "embedding_json": json.dumps(embedding) if embedding else "",
                        "token_count": len(chunk_text.split()),
                    }
                )

            self.write(
                {
                    "status": "indexed",
                    "last_indexed": fields.Datetime.now(),
                    "index_error": False,
                }
            )
        except Exception as exc:
            _logger.error("Indexing failed for document %d: %s", self.id, exc)
            self.write({"status": "error", "index_error": str(exc)[:500]})

    def _extract_text(self) -> str:
        """Extract plain text from content field or uploaded attachment.

        Supported formats: TXT, HTML, CSV (stdlib), PDF (pdfminer.six), DOCX (python-docx).
        PDF and DOCX fall back gracefully when the optional library is not installed.
        """
        if self.content:
            return self.content
        if not self.attachment_id or not self.attachment_id.datas:
            return ""

        import base64

        raw_bytes = base64.b64decode(self.attachment_id.datas)
        mime = (self.attachment_id.mimetype or "").lower()
        fname = (self.attachment_id.name or "").lower()

        # ── Plain text ─────────────────────────────────────────────────────────
        if "text/plain" in mime or fname.endswith(".txt"):
            return raw_bytes.decode("utf-8", errors="replace")

        # ── HTML ───────────────────────────────────────────────────────────────
        if "html" in mime or fname.endswith((".html", ".htm")):
            from html.parser import HTMLParser

            class _TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self._buf: list[str] = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style"):
                        self._skip = False

                def handle_data(self, data):
                    if not self._skip:
                        self._buf.append(data)

            parser = _TextExtractor()
            parser.feed(raw_bytes.decode("utf-8", errors="replace"))
            return " ".join(parser._buf)

        # ── CSV ────────────────────────────────────────────────────────────────
        if "csv" in mime or fname.endswith(".csv"):
            import csv
            import io

            text_io = io.StringIO(raw_bytes.decode("utf-8", errors="replace"))
            rows = list(csv.reader(text_io))
            return "\n".join(", ".join(row) for row in rows)

        # ── PDF ────────────────────────────────────────────────────────────────
        if "pdf" in mime or fname.endswith(".pdf"):
            try:
                import io

                from pdfminer.high_level import extract_text as pdf_extract_text

                return pdf_extract_text(io.BytesIO(raw_bytes)) or ""
            except ImportError:
                _logger.warning(
                    "pdfminer.six not installed — cannot extract text from PDF '%s'. "
                    "Install with: pip install pdfminer.six",
                    self.attachment_id.name,
                )
                return ""
            except Exception as exc:
                _logger.warning("PDF extraction failed for '%s': %s", self.attachment_id.name, exc)
                return ""

        # ── DOCX ───────────────────────────────────────────────────────────────
        if "docx" in mime or "openxmlformats" in mime or fname.endswith(".docx"):
            try:
                import io

                import docx

                doc = docx.Document(io.BytesIO(raw_bytes))
                return "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                _logger.warning(
                    "python-docx not installed — cannot extract text from DOCX '%s'. "
                    "Install with: pip install python-docx",
                    self.attachment_id.name,
                )
                return ""
            except Exception as exc:
                _logger.warning("DOCX extraction failed for '%s': %s", self.attachment_id.name, exc)
                return ""

        # ── Fallback: try UTF-8 decode ─────────────────────────────────────────
        try:
            return raw_bytes.decode("utf-8", errors="replace")
        except Exception:
            return ""

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
        """Split text into overlapping word-based chunks."""
        words = text.split()
        if not words:
            return []
        chunks = []
        start = 0
        while start < len(words):
            end = min(start + chunk_size, len(words))
            chunks.append(" ".join(words[start:end]))
            start += chunk_size - overlap
        return chunks


class AiDocumentChunk(models.Model):
    _name = "ai.document.chunk"
    _description = "AI Document Chunk"
    _order = "document_id, chunk_index"

    document_id = fields.Many2one("ai.document", required=True, ondelete="cascade", index=True)
    company_id = fields.Many2one("res.company", required=True, index=True)
    chunk_index = fields.Integer(readonly=True)
    content = fields.Text(readonly=True)
    token_count = fields.Integer(readonly=True)
    # Embedding stored as JSON (pgvector support deferred until extension is installed)
    embedding_json = fields.Text("Embedding (JSON)", readonly=True)

    @api.model
    def search_similar(
        self,
        query: str,
        company_id: int | None = None,
        limit: int = 5,
    ) -> list[dict]:
        """Return the *limit* most relevant chunks for *query*.

        Uses cosine similarity when pgvector is available; falls back to keyword
        search (ILIKE) otherwise.
        """
        cid = company_id or self.env.company.id

        if _check_pgvector(self.env):
            return self._vector_search(query, cid, limit)
        return self._keyword_search(query, cid, limit)

    # Common English/Dutch question & stop words that should not drive retrieval.
    _STOPWORDS = frozenset(
        {
            "what",
            "which",
            "when",
            "where",
            "who",
            "whom",
            "whose",
            "why",
            "how",
            "the",
            "a",
            "an",
            "is",
            "are",
            "was",
            "were",
            "be",
            "been",
            "being",
            "do",
            "does",
            "did",
            "to",
            "of",
            "in",
            "on",
            "at",
            "for",
            "and",
            "or",
            "with",
            "about",
            "into",
            "from",
            "that",
            "this",
            "these",
            "those",
            "wat",
            "welke",
            "wanneer",
            "waar",
            "wie",
            "hoe",
            "de",
            "het",
            "een",
            "zijn",
            "waren",
            "van",
            "voor",
            "met",
            "over",
        }
    )

    def _keyword_search(self, query: str, company_id: int, limit: int) -> list[dict]:
        """Fallback: ILIKE search on chunk content.

        Uses OR-matching across content words (stopwords/question words removed)
        and ranks results by how many distinct query terms each chunk contains,
        so natural-language questions like "What is the refund policy?" still
        retrieve the relevant chunk. (The old AND-all-terms logic returned
        nothing for questions because every chunk had to contain "what" too.)
        """
        raw_terms = [t.strip(".,?!;:\"'()").lower() for t in query.split()]
        terms = [t for t in raw_terms if len(t) > 2 and t not in self._STOPWORDS]
        if not terms:
            # Fall back to any token >2 chars if everything was a stopword
            terms = [t for t in raw_terms if len(t) > 2]
        if not terms:
            return []

        # OR across terms; rank by number of matched terms (relevance).
        or_clause = " OR ".join("content ILIKE %s" for _ in terms)
        score_expr = " + ".join("(content ILIKE %s)::int" for _ in terms)
        like_params = [f"%{t}%" for t in terms]
        params = like_params + like_params + [company_id, limit]
        self.env.cr.execute(
            f"""
            SELECT id, content, document_id, ({score_expr}) AS match_count
            FROM ai_document_chunk
            WHERE ({or_clause})
              AND company_id = %s
            ORDER BY match_count DESC
            LIMIT %s
            """,
            params,
        )
        rows = self.env.cr.fetchall()
        result = []
        n_terms = len(terms) or 1
        for row_id, content, doc_id, match_count in rows:
            doc = self.env["ai.document"].browse(doc_id)
            # Confidence scales with fraction of query terms matched (0.3–0.9).
            confidence = round(0.3 + 0.6 * (match_count / n_terms), 4)
            result.append(
                {
                    "chunk_id": row_id,
                    "content": content,
                    "source": doc.name,
                    "confidence": confidence,
                }
            )
        return result

    def _vector_search(self, query: str, company_id: int, limit: int) -> list[dict]:
        """pgvector cosine similarity search (requires vector extension + embeddings)."""
        provider = self.env["ai.provider"].get_default_provider()
        instance = provider.get_provider_instance() if provider else None
        if not instance:
            return self._keyword_search(query, company_id, limit)

        embedding = instance.embed(query)
        if not embedding:
            return self._keyword_search(query, company_id, limit)

        vec_str = "[" + ",".join(str(v) for v in embedding) + "]"
        try:
            self.env.cr.execute(
                """
                SELECT id, content, document_id,
                       1 - (embedding_json::vector <=> %s::vector) AS score
                FROM ai_document_chunk
                WHERE company_id = %s
                  AND embedding_json IS NOT NULL
                  AND embedding_json != ''
                ORDER BY embedding_json::vector <=> %s::vector
                LIMIT %s
                """,
                (vec_str, company_id, vec_str, limit),
            )
            rows = self.env.cr.fetchall()
            if not rows:
                return self._keyword_search(query, company_id, limit)
            result = []
            for row_id, content, doc_id, score in rows:
                doc = self.env["ai.document"].browse(doc_id)
                result.append(
                    {
                        "chunk_id": row_id,
                        "content": content,
                        "source": doc.name,
                        "confidence": round(float(score), 4),
                    }
                )
            return result
        except Exception as exc:
            _logger.warning("Vector search failed, falling back to keyword: %s", exc)
            return self._keyword_search(query, company_id, limit)
